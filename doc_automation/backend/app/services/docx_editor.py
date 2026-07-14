"""
services/docx_editor.py
────────────────────────
Core DOCX editing engine using python-docx.

Replaces placeholders in the Word template while preserving run formatting.
"""

from __future__ import annotations

from pathlib import Path
from docx import Document
from app.core.logging import get_logger
from app.models.schemas import ProjectData
from app.services.field_mapper import FieldMapper

log = get_logger(__name__)


class DocxEditor:
    """
    Stateless DOCX template editor using python-docx.
    Replaces matching placeholders in paragraphs and tables.
    """

    def __init__(self) -> None:
        self._mapper = FieldMapper()

    def edit_template(
        self,
        template_path: Path,
        output_path:   Path,
        project:       ProjectData,
        template_type: str = "plan",   # "plan" | "incubation"
    ) -> Path:
        """
        Edit *template_path* by replacing all matching placeholders
        with project data, then save to *output_path*.
        """
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Build replacement map from project data
        replacements = self._mapper.build_replacements(project, template_type)
        log.debug("[%s] %d replacement fields prepared", project.project_name, len(replacements))

        # Open the template
        doc = Document(str(template_path))

        # Perform replacement
        for placeholder, replacement in replacements.items():
            self._replace_placeholder(doc, placeholder, replacement)

        # Save the edited document
        doc.save(str(output_path))
        log.info("[%s] Saved → %s", project.project_name, output_path.name)
        return output_path

    def _replace_placeholder(self, doc: Document, placeholder: str, replacement: str) -> None:
        """Find and replace placeholder text in paragraphs and tables."""
        # 1. Process paragraphs in main document body
        for p in doc.paragraphs:
            if placeholder in p.text:
                self._replace_in_paragraph(p, placeholder, replacement)
            # Support smart normalization (if newlines are formatted as spaces)
            elif placeholder.replace("\n", " ") in p.text:
                self._replace_in_paragraph(p, placeholder.replace("\n", " "), replacement)
            elif placeholder.replace("\n", "") in p.text:
                self._replace_in_paragraph(p, placeholder.replace("\n", ""), replacement)

        # 2. Process paragraphs in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            self._replace_in_paragraph(p, placeholder, replacement)
                        elif placeholder.replace("\n", " ") in p.text:
                            self._replace_in_paragraph(p, placeholder.replace("\n", " "), replacement)
                        elif placeholder.replace("\n", "") in p.text:
                            self._replace_in_paragraph(p, placeholder.replace("\n", ""), replacement)

    def _replace_in_paragraph(self, paragraph, placeholder: str, replacement: str) -> None:
        """Replaces placeholder in a paragraph while trying to preserve formatting."""
        # Case 1: The placeholder is fully contained in a single run
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)
                return

        # Case 2: The placeholder is split across multiple runs
        text = paragraph.text
        start_idx = text.find(placeholder)
        if start_idx == -1:
            return
        end_idx = start_idx + len(placeholder)

        # Map character indices to run indices
        run_ranges = []
        current_char = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run_ranges.append((current_char, current_char + run_len))
            current_char += run_len

        # Find overlapping runs
        overlapping = []
        for idx, (r_start, r_end) in enumerate(run_ranges):
            if max(r_start, start_idx) < min(r_end, end_idx):
                overlapping.append(idx)

        if not overlapping:
            return

        first_run_idx = overlapping[0]
        last_run_idx = overlapping[-1]

        first_run = paragraph.runs[first_run_idx]
        last_run = paragraph.runs[last_run_idx]

        first_r_start, _ = run_ranges[first_run_idx]
        last_r_start, _ = run_ranges[last_run_idx]

        # Extract prefix from the first overlapping run
        prefix = first_run.text[:start_idx - first_r_start]
        # Extract suffix from the last overlapping run
        suffix = last_run.text[end_idx - last_r_start:]

        # Place prefix and replacement in first run
        first_run.text = prefix + replacement

        # Clear text of intermediate runs and last run
        for idx in overlapping[1:]:
            paragraph.runs[idx].text = ""

        # Place suffix in the last run (if it's different from the first)
        if first_run_idx != last_run_idx and suffix:
            last_run.text = suffix
